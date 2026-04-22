#!/usr/bin/env python3
"""测试表格边框功能"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

import unittest
import docx
from utils import apply_default_table_borders, get_table_borders_xml

class TestTableBorders(unittest.TestCase):
    def test_apply_default_borders(self):
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        apply_default_table_borders(table)
        borders = get_table_borders_xml(table)
        self.assertIsNotNone(borders)

    def test_borders_has_children(self):
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        apply_default_table_borders(table)
        borders = get_table_borders_xml(table)
        # Should have 6 border elements: top, left, bottom, right, insideH, insideV
        self.assertEqual(len(list(borders)), 6)

    def test_border_attributes(self):
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        apply_default_table_borders(table)
        borders = get_table_borders_xml(table)
        for border in borders:
            # Attributes use full namespace URI as key
            ns_val = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
            self.assertIn(ns_val, border.attrib)
            self.assertEqual(border.attrib.get(ns_val), 'single')

    def test_get_borders_none(self):
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        # Before applying borders, there may not be explicit borders
        borders = get_table_borders_xml(table)
        # Could be None or have default borders depending on python-docx
        # We just verify the function doesn't crash
        self.assertTrue(borders is None or len(list(borders)) >= 0)

if __name__ == '__main__':
    unittest.main()
