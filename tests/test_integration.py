#!/usr/bin/env python3
"""集成测试：验证完整流程"""
import sys
import os
import tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

import unittest
import docx
from docx.shared import Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from utils import (
    set_run_font, format_para_runs, get_para_props,
    apply_default_table_borders, fix_quotes,
    has_toc_field, emu_to_pt, pt_to_emu
)


class TestIntegration(unittest.TestCase):
    """集成测试：创建模板和目标文档，运行分析和修复操作"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.template_path = os.path.join(self.tmpdir, 'template.docx')
        self.target_path = os.path.join(self.tmpdir, 'target.docx')

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _create_template(self):
        """创建一个简单的模板文档"""
        doc = docx.Document()
        # 封面标题
        p = doc.add_paragraph('《测试文档》')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, '宋体', 304800, True)

        # 正文标题
        p = doc.add_paragraph('一、测试章节')
        p.paragraph_format.first_line_indent = Emu(356870)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            set_run_font(run, '黑体', 177800, True)

        # 正文段落
        p = doc.add_paragraph('这是一段正文内容。')
        p.paragraph_format.first_line_indent = Emu(304800)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            set_run_font(run, '宋体', 152400, False)

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        apply_default_table_borders(table)
        table.cell(0, 0).text = '表头1'
        table.cell(0, 1).text = '表头2'
        table.cell(1, 0).text = '数据1'
        table.cell(1, 1).text = '数据2'

        doc.save(self.template_path)
        return doc

    def _create_target(self):
        """创建一个格式混乱的目标文档"""
        doc = docx.Document()
        # 封面标题（不同格式）
        p = doc.add_paragraph('《测试文档》')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 错误对齐
        for run in p.runs:
            set_run_font(run, '微软雅黑', 200000, False)  # 错误字体字号

        # 正文标题（不同格式）
        p = doc.add_paragraph('一、测试章节')
        p.paragraph_format.first_line_indent = None  # 缺少缩进
        p.paragraph_format.line_spacing = 1.0  # 错误行距
        for run in p.runs:
            set_run_font(run, '楷体', 120000, False)  # 错误字体字号

        # 正文段落（包含半角引号）
        p = doc.add_paragraph('这是一段包含"引号"的正文内容。')
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            set_run_font(run, '微软雅黑', 120000, False)

        # 添加无边框表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '表头1'
        table.cell(0, 1).text = '表头2'

        doc.save(self.target_path)
        return doc

    def test_create_documents(self):
        """测试：能够创建模板和目标文档"""
        tmpl = self._create_template()
        target = self._create_target()
        self.assertTrue(os.path.exists(self.template_path))
        self.assertTrue(os.path.exists(self.target_path))

    def test_analyze_template(self):
        """测试：能够分析模板文件"""
        self._create_template()
        doc = docx.Document(self.template_path)

        # 验证段落属性提取
        props_list = []
        for p in doc.paragraphs:
            props = get_para_props(p)
            if props:
                props_list.append(props)

        self.assertEqual(len(props_list), 3)
        self.assertEqual(props_list[0]['text'], '《测试文档》')
        self.assertEqual(props_list[0]['alignment'], WD_ALIGN_PARAGRAPH.CENTER)

    def test_fix_quotes_on_target(self):
        """测试：能够修复目标文档中的引号"""
        self._create_target()
        doc = docx.Document(self.target_path)

        # 修复前包含半角引号
        has_halfwidth = any('"' in p.text for p in doc.paragraphs)
        self.assertTrue(has_halfwidth)

        fix_quotes(doc)

        # 修复后不应包含半角引号
        has_halfwidth = any('"' in p.text for p in doc.paragraphs)
        self.assertFalse(has_halfwidth)

    def test_apply_table_borders(self):
        """测试：能够给目标文档表格添加边框"""
        self._create_target()
        doc = docx.Document(self.target_path)

        table = doc.tables[0]
        apply_default_table_borders(table)

        # 验证边框已添加
        from utils import get_table_borders_xml
        borders = get_table_borders_xml(table)
        self.assertIsNotNone(borders)
        self.assertEqual(len(list(borders)), 6)

    def test_para_props_comparison(self):
        """测试：能够比较模板和目标的段落属性差异"""
        self._create_template()
        self._create_target()

        tmpl = docx.Document(self.template_path)
        target = docx.Document(self.target_path)

        # 提取模板属性
        tmpl_map = {}
        for p in tmpl.paragraphs:
            info = get_para_props(p)
            if info:
                tmpl_map[info['text']] = info

        # 提取目标属性
        target_map = {}
        for p in target.paragraphs:
            info = get_para_props(p)
            if info:
                target_map[info['text']] = info

        # 验证差异
        self.assertIn('《测试文档》', tmpl_map)
        self.assertIn('《测试文档》', target_map)

        # 对齐方式不同
        self.assertEqual(tmpl_map['《测试文档》']['alignment'], WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(target_map['《测试文档》']['alignment'], WD_ALIGN_PARAGRAPH.LEFT)

    def test_import_all_scripts(self):
        """测试：所有脚本模块都可以被导入"""
        scripts_dir = os.path.join(os.path.dirname(__file__), '..', 'scripts')

        # 验证关键模块可导入
        try:
            import utils
            self.assertTrue(hasattr(utils, 'emu_to_pt'))
            self.assertTrue(hasattr(utils, 'fix_quotes'))
            self.assertTrue(hasattr(utils, 'apply_default_table_borders'))
        except ImportError as e:
            self.fail(f'导入 utils 失败: {e}')

        # 验证脚本文件存在
        expected_files = [
            'analyze_template.py',
            'audit_docx.py',
            'fix_docx_template.py',
            'verify_docx.py',
            'copy_styles.py',
            'copy_headers_footers.py',
            'utils.py',
        ]
        for f in expected_files:
            path = os.path.join(scripts_dir, f)
            self.assertTrue(os.path.exists(path), f'缺少文件: {f}')

    def test_emu_conversions_roundtrip(self):
        """测试：EMU 换算往返一致性"""
        test_values = [1, 10, 12, 24, 72]
        for pt in test_values:
            emu = pt_to_emu(pt)
            pt_back = emu_to_pt(emu)
            self.assertAlmostEqual(pt, pt_back, places=5)


if __name__ == '__main__':
    unittest.main()
