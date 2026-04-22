#!/usr/bin/env python3
"""测试引号转换功能"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

import unittest
import docx
from utils import fix_quotes

class TestFixQuotes(unittest.TestCase):
    def test_halfwidth_to_fullwidth_double(self):
        doc = docx.Document()
        p = doc.add_paragraph('他说:"你好。"')
        fix_quotes(doc)
        self.assertIn('\u201c', p.text)  # "
        self.assertIn('\u201d', p.text)  # "
        self.assertNotIn('"', p.text)    # original "

    def test_single_quotes(self):
        doc = docx.Document()
        p = doc.add_paragraph("It's a 'test'.")
        fix_quotes(doc)
        # Should convert single quotes too
        self.assertNotIn("'", p.text)

    def test_chinese_quotes(self):
        doc = docx.Document()
        p = doc.add_paragraph('"这是一个测试"')
        fix_quotes(doc)
        self.assertIn('\u201c', p.text)
        self.assertIn('\u201d', p.text)
        self.assertNotIn('"', p.text)

    def test_mixed_quotes(self):
        doc = docx.Document()
        p = doc.add_paragraph('"他说:\'你好\'"')
        fix_quotes(doc)
        self.assertNotIn('"', p.text)
        self.assertNotIn("'", p.text)

    def test_no_quotes(self):
        doc = docx.Document()
        p = doc.add_paragraph('这是一段没有引号的文字。')
        fix_quotes(doc)
        self.assertEqual(p.text, '这是一段没有引号的文字。')

if __name__ == '__main__':
    unittest.main()
