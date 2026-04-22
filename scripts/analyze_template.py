#!/usr/bin/env python3
"""
深度分析模板文件格式规范 - 输出显式格式和隐藏格式
包含：段落、分节、页眉页脚、样式、表格、目录域
"""
import docx
from docx.shared import Emu, Inches
from docx.enum.section import WD_ORIENT
import sys
import re
import argparse

from utils import emu_to_pt, emu_to_inch, get_para_props, ns_tag, ns_attr
from utils import run_with_errors, check_write_permission, log_ok, log_warn, log_err

parser = argparse.ArgumentParser(description='深度分析模板文件格式规范')
parser.add_argument('template', nargs='?', default='template.docx', help='模板 docx 文件路径')
args = parser.parse_args()
TEMPLATE = args.template

# ============== 段落分析 ==============
def get_para_props_with_idx(p, idx):
    info = get_para_props(p, max_text_len=70, with_outline=True)
    if info is None:
        return None
    info['idx'] = idx
    return info

def classify(text, idx):
    """根据内容分类段落类型"""
    if idx == 0 and text.startswith('《') and text.endswith('》'):
        return '封面标题'
    if text == '课程标准':
        return '"课程标准"'
    if text == '目录':
        return '"目录"'
    if '课程标准' in text and text.startswith('《') and idx > 10:
        return '正文大标题'
    if re.match(r'^[一二三四五六七八九十]+、', text) and '\t' in text:
        return '目录项_toc1'
    if re.match(r'^（[一二三四五六七八九十]+）', text) and '\t' in text:
        return '目录项_toc2'
    if text.startswith('附录：\t'):
        return '目录项_toc1'
    if re.match(r'^[一二三]、课程', text) and '\t' not in text and len(text) < 20:
        return '一级标题_一二三'
    if re.match(r'^[四五六]、', text) and '\t' not in text and len(text) < 20:
        return '一级标题_四五六'
    if re.match(r'^（[一二三四五六七八九十]+）', text) and '\t' not in text and len(text) < 20:
        return '二级标题'
    if text in ['评价方式', '评价比例']:
        return '评价标题'
    if re.match(r'^表[123456789]', text):
        return '表标题'
    if text == '附录：' and idx > 50 and '\t' not in text:
        return '附录正文'
    if re.match(r'^[123]\.(知识|能力|素质)目标', text):
        return '目标标题'
    if re.match(r'^[12]\.\s+(校内|企业)', text):
        return '师资标题'
    if re.match(r'^[123]\.(\s+|．)(文本|校内|地域)', text):
        return '资源标题'
    if re.match(r'^[123]\)[^\)]', text):
        return '编号要求'
    if re.match(r'^[1234]\)\.\s+', text):
        return '编号列表'
    if re.match(r'^\([123456789]\)', text):
        return '括号编号'
    if re.match(r'^（[1234]）', text) and len(text) > 5:
        return '括号中文编号正文'
    return '正文'

# ============== 分节/页面设置分析 ==============
def analyze_sections(doc):
    print("\n" + "=" * 100)
    print("【分节与页面设置】")
    print("=" * 100)

    for i, section in enumerate(doc.sections):
        print(f"\n--- Section {i+1} ---")
        print(f"  纸张宽度: {emu_to_inch(section.page_width):.3f}in ({section.page_width})")
        print(f"  纸张高度: {emu_to_inch(section.page_height):.3f}in ({section.page_height})")
        print(f"  方向: {'横向' if section.orientation == WD_ORIENT.LANDSCAPE else '纵向'}")
        print(f"  上边距: {emu_to_inch(section.top_margin):.3f}in ({section.top_margin})")
        print(f"  下边距: {emu_to_inch(section.bottom_margin):.3f}in ({section.bottom_margin})")
        print(f"  左边距: {emu_to_inch(section.left_margin):.3f}in ({section.left_margin})")
        print(f"  右边距: {emu_to_inch(section.right_margin):.3f}in ({section.right_margin})")
        print(f"  页眉距边界: {emu_to_inch(section.header_distance):.3f}in ({section.header_distance})")
        print(f"  页脚距边界: {emu_to_inch(section.footer_distance):.3f}in ({section.footer_distance})")
        print(f"  首页不同: {section.different_first_page_header_footer}")
        # 奇偶页不同需检查 sectPr 中的 evenAndOddHeaders 元素
        sectPr = section._sectPr
        even_odd = sectPr.find(ns_tag('w:evenAndOddHeaders'))
        print(f"  奇偶页不同: {even_odd is not None}")

        # 分节符类型
        sectPr = section._sectPr
        sect_type = sectPr.find(ns_tag('w:type'))
        if sect_type is not None:
            print(f"  分节符类型: {sect_type.get(ns_attr('val'), 'default')}")

# ============== 页眉页脚分析 ==============
def analyze_headers_footers(doc):
    print("\n" + "=" * 100)
    print("【页眉页脚】")
    print("=" * 100)

    for i, section in enumerate(doc.sections):
        print(f"\n--- Section {i+1} ---")

        # 页眉
        header = section.header
        if header and header.paragraphs:
            print(f"  页眉段落数: {len(header.paragraphs)}")
            for j, p in enumerate(header.paragraphs):
                text = p.text.strip()
                if text:
                    runs = p.runs
                    first_run = runs[0] if runs else None
                    print(f"    [{j}] '{text[:50]}' font={first_run.font.name if first_run else None} size={first_run.font.size if first_run else None} bold={first_run.font.bold if first_run else None}")
        else:
            print("  页眉: 无")

        # 页脚
        footer = section.footer
        if footer and footer.paragraphs:
            print(f"  页脚段落数: {len(footer.paragraphs)}")
            for j, p in enumerate(footer.paragraphs):
                text = p.text.strip()
                if text:
                    runs = p.runs
                    first_run = runs[0] if runs else None
                    print(f"    [{j}] '{text[:50]}' font={first_run.font.name if first_run else None} size={first_run.font.size if first_run else None} bold={first_run.font.bold if first_run else None}")

                    # 检测域代码
                    for run in p.runs:
                        if run._element.findall(f'.//{ns_tag("w:fldChar")}'):
                            print(f"        -> 包含域代码")
        else:
            print("  页脚: 无")

# ============== 样式分析 ==============
def analyze_styles(doc):
    print("\n" + "=" * 100)
    print("【样式定义】")
    print("=" * 100)

    for style in doc.styles:
        if style.type is None:
            continue

        # 只输出关键样式
        key_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                      'toc 1', 'toc 2', 'toc 3', 'Header', 'Footer',
                      'Title', 'Subtitle']
        if style.name in key_styles or style.type == 1:  # 1 = paragraph style
            print(f"\n  样式名: {style.name} (type={style.type})")

            # 字体
            if style.font:
                print(f"    font_name={style.font.name}, size={style.font.size}, bold={style.font.bold}")

            # 段落格式
            pf = style.paragraph_format
            if pf:
                print(f"    alignment={pf.alignment}, line_spacing={pf.line_spacing}")
                print(f"    space_before={pf.space_before}, space_after={pf.space_after}")
                print(f"    first_line_indent={pf.first_line_indent}, left_indent={pf.left_indent}")

# ============== 表格分析 ==============
def analyze_tables(doc):
    print("\n" + "=" * 100)
    print("【表格】")
    print("=" * 100)

    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i+1} ---")
        print(f"  行数: {len(table.rows)}, 列数: {len(table.columns)}")

        # 表格边框
        tbl = table._tbl
        tblPr = tbl.find(ns_tag('w:tblPr'))
        if tblPr is not None:
            borders = tblPr.find(ns_tag('w:tblBorders'))
            if borders is not None:
                for border in borders:
                    tag = border.tag.split('}')[-1]
                    val = border.get(ns_tag('w:val'))
                    sz = border.get(ns_tag('w:sz'))
                    print(f"  边框 {tag}: val={val}, sz={sz}")
            else:
                print("  边框: 无显式设置（继承默认）")

        # 表格宽度
        tblW = tblPr.find(ns_tag('w:tblW')) if tblPr is not None else None
        if tblW is not None:
            w_type = tblW.get(ns_tag('w:type'))
            w_val = tblW.get(ns_tag('w:w'))
            print(f"  表格宽度: type={w_type}, w={w_val}")

        # 对齐方式
        jc = tblPr.find(ns_tag('w:jc')) if tblPr is not None else None
        if jc is not None:
            print(f"  表格对齐: {jc.get(ns_tag('w:val'))}")

        # 列宽
        print("  列宽:")
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                tcPr = cell._tc.find(ns_tag('w:tcPr'))
                if tcPr is not None:
                    tcW = tcPr.find(ns_tag('w:tcW'))
                    if tcW is not None:
                        w_type = tcW.get(ns_tag('w:type'))
                        w_val = tcW.get(ns_tag('w:w'))
                        print(f"    col{j}: type={w_type}, w={w_val}")
            break  # 只输出第一行

        # 第一行内容（标题行）
        print("  第一行内容:")
        if table.rows:
            for j, cell in enumerate(table.rows[0].cells):
                text = cell.text.strip()[:30]
                print(f"    [{j}] '{text}'")

# ============== 目录域分析 ==============
def analyze_toc(doc):
    print("\n" + "=" * 100)
    print("【目录域代码】")
    print("=" * 100)

    found_toc = False

    for i, p in enumerate(doc.paragraphs):
        # 查找 fldSimple 域
        fld_simple = p._element.find(f'.//{ns_tag("w:fldSimple")}')
        if fld_simple is not None:
            instr = fld_simple.get(ns_attr('instr'))
            if instr and 'TOC' in instr:
                found_toc = True
                print(f"\n  [{i}] 发现 TOC 域 (fldSimple):")
                print(f"    指令: {instr}")

        # 查找复杂域 (fldChar)
        fld_chars = p._element.findall(f'.//{ns_tag("w:fldChar")}')
        if fld_chars:
            for fld in fld_chars:
                fld_type = fld.get(ns_attr('fldCharType'))
                if fld_type == 'begin':
                    # 查找对应的 instrText
                    instr_texts = p._element.findall(f'.//{ns_tag("w:instrText")}')
                    for it in instr_texts:
                        if it.text and 'TOC' in it.text:
                            found_toc = True
                            print(f"\n  [{i}] 发现 TOC 域 (fldChar):")
                            print(f"    指令: {it.text}")

    if not found_toc:
        print("  未发现 TOC 域（可能是手动目录或静态文本）")

# ============== 主程序 ==============
@run_with_errors
def main():
    doc = docx.Document(TEMPLATE)

    print("=" * 100)
    print(f"深度分析模板文件: {TEMPLATE}")
    print("=" * 100)

    # 1. 段落分析
    print("\n" + "=" * 100)
    print("【段落格式分析】")
    print("=" * 100)

    results = {}
    for i, p in enumerate(doc.paragraphs):
        info = get_para_props_with_idx(p, i)
        if not info:
            continue
        category = classify(info['text'], i)
        if category not in results:
            results[category] = []
        results[category].append(info)

    for cat, items in sorted(results.items()):
        print(f"\n【{cat}】出现 {len(items)} 次，示例索引: {[x['idx'] for x in items[:3]]}")
        first = items[0]
        print(f"  示例: '{first['text'][:45]}'")
        print(f"  style={first['style']}, outline_lvl={first['outline_lvl']}")
        print(f"  font={first['font_name']}, size={first['font_size']}({emu_to_pt(first['font_size'])}pt), bold={first['bold']}")
        print(f"  alignment={first['alignment']}, line_spacing={first['line_spacing']}")
        print(f"  sb={first['space_before']}, sa={first['space_after']}, fi={first['first_line_indent']}")

    # 2. 分节分析
    analyze_sections(doc)

    # 3. 页眉页脚
    analyze_headers_footers(doc)

    # 4. 样式
    analyze_styles(doc)

    # 5. 表格
    analyze_tables(doc)

    # 6. 目录
    analyze_toc(doc)

    # 汇总
    print("\n" + "=" * 100)
    print("【格式规范汇总表】")
    print("=" * 100)
    print(f"{'类型':<20} {'样式':<15} {'字体':<8} {'字号':<10} {'加粗':<6} {'斜体':<6} {'下划线':<6} {'对齐':<10} {'行距':<12} {'段前':<10} {'段后':<10} {'首行缩进':<10}")
    print("-" * 135)
    for cat, items in sorted(results.items()):
        first = items[0]
        size_pt = emu_to_pt(first['font_size']) if first['font_size'] else None
        print(f"{cat:<20} {str(first['style']):<15} {str(first['font_name']):<8} {str(size_pt)+'pt':<10} {str(first['bold']):<6} {str(first['italic']):<6} {str(first['underline']):<6} {str(first['alignment']):<10} {str(first['line_spacing'])[:11]:<12} {str(first['space_before']):<10} {str(first['space_after']):<10} {str(first['first_line_indent']):<10}")

    print("\n" + "=" * 100)
    print(f"分析完成。段落总数: {len(doc.paragraphs)}, 表格总数: {len(doc.tables)}, 分节总数: {len(doc.sections)}")
    print("=" * 100)

if __name__ == '__main__':
    main()
