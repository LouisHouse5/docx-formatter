#!/usr/bin/env python3
"""
修复 docx 文件格式 - 以模板为基准统一所有格式（显式+隐藏）
使用方法:
1. 先用 analyze_template.py 分析模板，获取精确参数
2. 修改下方 CONFIG 部分，填入对应参数
3. 修改 classify_and_format() 中的匹配规则（如有特殊段落类型）
4. 修改 fix_specific_issues() 处理文件特定问题
5. 如需同步样式/页眉页脚，先运行 copy_styles.py / copy_headers_footers.py
6. 运行: python3 fix_docx.py
"""
import docx
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re
import copy

# ========================== CONFIG 配置区（根据模板修改） ==========================
TARGET = '目标文件.docx'  # <-- 修改为目标文件名
TEMPLATE = 'template.docx'  # <-- 修改为模板文件名（用于自动同步隐藏格式）

# 字体配置
FONT_SONG = '宋体'
FONT_HEI = '黑体'

# 字号（EMU）
SIZE_COVER_TITLE = 304800      # 封面标题
SIZE_COVER_SUBTITLE = 330200   # "课程标准"
SIZE_TOC = 152400              # 目录项
SIZE_MAIN_TITLE = 203200       # 正文大标题
SIZE_H1 = 177800               # 一级标题
SIZE_H2 = 152400               # 二级标题
SIZE_BODY = 152400             # 正文
SIZE_TABLE_TITLE = 152400      # 表标题

# 首行缩进（EMU）
FI_H1 = 356870          # 一级标题（一、二、三）
FI_H1_2 = 266700        # 一级标题（四、五、六）
FI_H2 = 306070          # 二级标题
FI_BODY = 304800        # 正文
FI_REQUIRE = 300355     # 1)2)3) 要求
FI_RESOURCE = 459105    # 资源标题

# 段前段后（EMU）
SB_COVER_TITLE = 1188720
SA_COVER_TITLE = 396240
SA_COVER_SUBTITLE = 3566160
SB_H1 = 76200
SB_H1_2 = 228600
SA_MAIN_TITLE = 99060

# 行距
LS_COVER = Pt(3)
LS_NORMAL = 1.5
LS_TOC = 2.0
LS_H1_FIRST = 2.5

# 页面设置（如需同步）
PAGE_WIDTH = Inches(8.27)       # A4 宽度
PAGE_HEIGHT = Inches(11.69)     # A4 高度
TOP_MARGIN = Inches(1.02)
BOTTOM_MARGIN = Inches(1.02)
LEFT_MARGIN = Inches(1.02)
RIGHT_MARGIN = Inches(1.02)

# =================================================================================

def set_run_font(run, name=FONT_SONG, size=SIZE_BODY, bold=False, italic=False, underline=False, color=None):
    """设置 run 字体，包含东亚字体属性和完整格式"""
    run.font.name = name
    run.font.size = Emu(size)
    if bold is None:
        run.font.bold = None
    else:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if underline is not None:
        run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def format_para_runs(p, name=FONT_SONG, size=SIZE_BODY, bold=False):
    """格式化段落中所有 run"""
    for run in p.runs:
        set_run_font(run, name, size, bold)

def insert_paragraph_before(paragraph, text):
    """在指定段落之前插入新段落（XML 级别）"""
    new_p = docx.oxml.OxmlElement('w:p')
    paragraph._element.addprevious(new_p)
    new_run = docx.oxml.OxmlElement('w:r')
    new_p.append(new_run)
    new_text = docx.oxml.OxmlElement('w:t')
    new_text.text = text
    new_run.append(new_text)
    return docx.text.paragraph.Paragraph(new_p, paragraph._parent)

def apply_style(p, doc, style_name):
    """应用指定样式，如样式不存在则忽略"""
    try:
        p.style = doc.styles[style_name]
    except:
        pass

def classify_and_format(doc):
    """
    遍历所有段落，根据内容类型匹配模板格式
    如需添加特殊段落类型，在此添加匹配规则
    """
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # ---- 1. 封面标题 ----
        if i == 0 and text.startswith('《') and text.endswith('》'):
            p.alignment = None
            p.paragraph_format.line_spacing = LS_COVER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p.paragraph_format.space_before = Emu(SB_COVER_TITLE)
            p.paragraph_format.space_after = Emu(SA_COVER_TITLE)
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_COVER_TITLE, True)

        # ---- 2. "课程标准" ----
        elif text == '课程标准':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = Emu(SA_COVER_SUBTITLE)
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_COVER_SUBTITLE, True)

        # ---- 3. "目录" ----
        elif text == '目录':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = None
            p.paragraph_format.line_spacing_rule = None
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_COVER_TITLE, True)

        # ---- 4. 正文大标题 ----
        elif '课程标准' in text and text.startswith('《') and i > 10:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = None
            p.paragraph_format.line_spacing_rule = None
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = Emu(SA_MAIN_TITLE)
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_MAIN_TITLE, True)

        # ---- 5. 目录项 toc1 ----
        elif re.match(r'^[一二三四五六七八九十]+、', text) and '\t' in text:
            apply_style(p, doc, 'toc 1')
            p.paragraph_format.line_spacing = LS_TOC
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = None
            is_level1 = bool(re.match(r'^[一二三四五六七八九十]+、[^（]', text))
            format_para_runs(p, FONT_SONG, SIZE_TOC, is_level1)

        # ---- 6. 目录项 toc2 ----
        elif re.match(r'^（[一二三四五六七八九十]+）', text) and '\t' in text:
            apply_style(p, doc, 'toc 2')
            p.paragraph_format.line_spacing = LS_TOC
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_TOC, False)

        # ---- 7. 附录目录项 ----
        elif text.startswith('附录：\t'):
            apply_style(p, doc, 'toc 1')
            p.paragraph_format.line_spacing = LS_TOC
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_TOC, True)

        # ---- 8. 一级标题（一、二、三）----
        elif re.match(r'^[一二三]、课程', text) and '\t' not in text and len(text) < 20:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_style(p, doc, 'Heading 1')
            if i > 30:
                p.paragraph_format.line_spacing = LS_NORMAL
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            else:
                p.paragraph_format.line_spacing = LS_H1_FIRST
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.space_before = Emu(SB_H1)
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_H1)
            format_para_runs(p, FONT_SONG, SIZE_H1, True)

        # ---- 9. 一级标题（四、五、六）----
        elif re.match(r'^[四五六]、', text) and '\t' not in text and len(text) < 20:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_style(p, doc, 'Heading 1')
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if text.startswith('六、'):
                p.paragraph_format.space_before = None
                p.paragraph_format.first_line_indent = Emu(FI_H1)
            else:
                p.paragraph_format.space_before = Emu(SB_H1_2)
                p.paragraph_format.first_line_indent = Emu(FI_H1_2)
            p.paragraph_format.space_after = None
            format_para_runs(p, FONT_SONG, SIZE_H1, True)

        # ---- 10. 二级标题 ----
        elif re.match(r'^（[一二三四五六七八九十]+）', text) and '\t' not in text and len(text) < 20:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_style(p, doc, 'Heading 2')
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_H2)
            format_para_runs(p, FONT_SONG, SIZE_H2, True)

        # ---- 11. 评价方式/评价比例 ----
        elif text in ['评价方式', '评价比例']:
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

        # ---- 12. 附录正文 ----
        elif text == '附录：' and i > 50 and '\t' not in text:
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_H1)
            format_para_runs(p, FONT_SONG, SIZE_H1, True)

        # ---- 13. 表标题 ----
        elif re.match(r'^表[123456789]', text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = Emu(99060) if '表1' in text or '表5' in text else None
            p.paragraph_format.space_after = Emu(99060) if '表2' in text or '表3' in text else None
            p.paragraph_format.first_line_indent = None
            format_para_runs(p, FONT_SONG, SIZE_TABLE_TITLE, True)

        # ---- 14. 1.2.3. 知识/能力/素质目标 ----
        elif re.match(r'^[123]\.(知识|能力|素质)目标', text):
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_H2)
            format_para_runs(p, FONT_SONG, SIZE_BODY, True)

        # ---- 15. (1)-(9) 列表 ----
        elif re.match(r'^\([123456789]\)', text):
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_BODY)
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

        # ---- 16. （1）-（4）师资条件 ----
        elif re.match(r'^（[1234]）', text) and len(text) > 5:
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_BODY)
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

        # ---- 17. 1. 校内/企业教师 ----
        elif re.match(r'^[12]\.\s+(校内|企业)', text):
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_H2)
            format_para_runs(p, FONT_SONG, SIZE_BODY, True)

        # ---- 18. 1. 文本/校内/地域资源 ----
        elif re.match(r'^[123]\.(\s+|．)(文本|校内|地域)', text):
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_RESOURCE)
            format_para_runs(p, FONT_SONG, SIZE_BODY, True)

        # ---- 19. 1). 2). 3). 4). 资源列表 ----
        elif re.match(r'^[1234]\)\.\s+', text):
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_BODY)
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

        # ---- 20. 1) 2) 3) 硬件/软件/其他要求 ----
        elif re.match(r'^[123]\)[^\)]', text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            p.paragraph_format.first_line_indent = Emu(FI_REQUIRE)
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

        # ---- 21. 默认正文 ----
        else:
            p.paragraph_format.line_spacing = LS_NORMAL
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = None
            p.paragraph_format.space_after = None
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                p.paragraph_format.first_line_indent = Emu(FI_BODY)
            format_para_runs(p, FONT_SONG, SIZE_BODY, False)

def fix_quotes(doc):
    """半角引号转全角引号（智能开闭匹配，包含单双引号）"""
    def smart_quote_replace(text):
        """智能判断开闭引号"""
        result = []
        for i, ch in enumerate(text):
            if ch == '"':
                # 段落开头、空格或中文标点前为开引号
                if i == 0 or text[i-1] in ' \t\n\r（【《「『"\u201c\u2018：。，？！；、':
                    result.append('\u201c')  # "
                else:
                    result.append('\u201d')  # "
            elif ch == "'":
                if i == 0 or text[i-1] in ' \t\n\r（【《「『\u201c\u2018：。，？！；、':
                    result.append('\u2018')  # '
                else:
                    result.append('\u2019')  # '
            else:
                result.append(ch)
        return ''.join(result)

    for p in doc.paragraphs:
        for run in p.runs:
            if '"' in run.text or "'" in run.text:
                run.text = smart_quote_replace(run.text)

def fix_table_fonts(doc):
    """统一表格内字体"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, FONT_SONG, SIZE_BODY, run.font.bold or False)

def _get_table_borders_xml(table):
    """提取表格的边框 XML 元素（深拷贝）"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tbl = table._tbl
    tblPr = tbl.find(f'{{{ns}}}tblPr')
    if tblPr is None:
        return None
    borders = tblPr.find(f'{{{ns}}}tblBorders')
    if borders is None:
        return None
    return copy.deepcopy(borders)

def _apply_table_borders(table, borders_element):
    """将边框 XML 应用到表格"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tbl = table._tbl
    tblPr = tbl.find(f'{{{ns}}}tblPr')
    if tblPr is None:
        tblPr = docx.oxml.OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    old_borders = tblPr.find(f'{{{ns}}}tblBorders')
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(copy.deepcopy(borders_element))

def copy_table_borders_from_template(tmpl_doc, target_doc):
    """从模板自动同步表格边框（按表格索引一对一）"""
    tmpl_tables = tmpl_doc.tables
    target_tables = target_doc.tables
    min_tables = min(len(tmpl_tables), len(target_tables))

    copied = 0
    for i in range(min_tables):
        borders = _get_table_borders_xml(tmpl_tables[i])
        if borders is not None:
            _apply_table_borders(target_tables[i], borders)
            copied += 1

    print(f"  已同步 {copied}/{len(target_tables)} 个表格边框")

def fix_table_borders(doc):
    """
    【兼容旧版】统一表格边框样式（XML 级别操作）
    当未提供模板时，使用硬编码默认边框
    """
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(f'{{{ns}}}tblPr')
        if tblPr is None:
            tblPr = docx.oxml.OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        borders = tblPr.find(f'{{{ns}}}tblBorders')
        if borders is None:
            borders = docx.oxml.OxmlElement('w:tblBorders')
            tblPr.append(borders)

        # 清除旧边框设置
        for child in list(borders):
            borders.remove(child)

        # 添加新边框（默认：所有边框单线，宽度 4 = 1/2 pt）
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = docx.oxml.OxmlElement(f'w:{border_name}')
            border.set(f'{{{ns}}}val', 'single')
            border.set(f'{{{ns}}}sz', '4')
            border.set(f'{{{ns}}}space', '0')
            border.set(f'{{{ns}}}color', 'auto')
            borders.append(border)

def fix_table_alignment(doc):
    """统一表格对齐方式"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(f'{{{ns}}}tblPr')
        if tblPr is None:
            tblPr = docx.oxml.OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        jc = tblPr.find(f'{{{ns}}}jc')
        if jc is None:
            jc = docx.oxml.OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(f'{{{ns}}}val', 'center')

def _get_odd_even_headers(section):
    """检查 section 是否启用奇偶页不同页眉页脚"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    sectPr = section._sectPr
    return sectPr.find(f'{{{ns}}}evenAndOddHeaders') is not None


def _set_odd_even_headers(section, enabled):
    """设置 section 的奇偶页不同页眉页脚"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    sectPr = section._sectPr
    existing = sectPr.find(f'{{{ns}}}evenAndOddHeaders')
    if enabled:
        if existing is None:
            sectPr.append(docx.oxml.OxmlElement('w:evenAndOddHeaders'))
    else:
        if existing is not None:
            sectPr.remove(existing)


def fix_section_settings(doc, tmpl_doc=None):
    """
    同步页面设置。
    如果提供了 tmpl_doc，按 section 索引从模板一对一复制；
    否则使用 CONFIG 中的硬编码值统一应用到所有 section。
    """
    if tmpl_doc is not None:
        tmpl_sections = tmpl_doc.sections
        target_sections = doc.sections
        min_sections = min(len(tmpl_sections), len(target_sections))

        for i in range(min_sections):
            ts = tmpl_sections[i]
            s = target_sections[i]
            s.page_width = ts.page_width
            s.page_height = ts.page_height
            s.orientation = ts.orientation
            s.top_margin = ts.top_margin
            s.bottom_margin = ts.bottom_margin
            s.left_margin = ts.left_margin
            s.right_margin = ts.right_margin
            s.header_distance = ts.header_distance
            s.footer_distance = ts.footer_distance
            s.different_first_page_header_footer = ts.different_first_page_header_footer
            # odd_and_even_pages_header_footer 不是 python-docx 属性，需操作 XML
            _set_odd_even_headers(s, _get_odd_even_headers(ts))

        if len(target_sections) > len(tmpl_sections):
            # 目标 section 更多，剩余的使用最后一个模板的设置
            last_tmpl = tmpl_sections[-1]
            for i in range(len(tmpl_sections), len(target_sections)):
                s = target_sections[i]
                s.page_width = last_tmpl.page_width
                s.page_height = last_tmpl.page_height
                s.orientation = last_tmpl.orientation
                s.top_margin = last_tmpl.top_margin
                s.bottom_margin = last_tmpl.bottom_margin
                s.left_margin = last_tmpl.left_margin
                s.right_margin = last_tmpl.right_margin
                s.header_distance = last_tmpl.header_distance
                s.footer_distance = last_tmpl.footer_distance
                s.different_first_page_header_footer = last_tmpl.different_first_page_header_footer
                _set_odd_even_headers(s, _get_odd_even_headers(last_tmpl))
        print(f"  已同步 {len(target_sections)} 个 section 的页面设置（来自模板）")
    else:
        for section in doc.sections:
            section.page_width = PAGE_WIDTH
            section.page_height = PAGE_HEIGHT
            section.top_margin = TOP_MARGIN
            section.bottom_margin = BOTTOM_MARGIN
            section.left_margin = LEFT_MARGIN
            section.right_margin = RIGHT_MARGIN
        print(f"  已同步 {len(doc.sections)} 个 section 的页面设置（来自 CONFIG）")

def remove_empty_table_rows(doc):
    """删除表格末尾空行"""
    for table in doc.tables:
        empty_rows = 0
        for row in reversed(table.rows):
            cells_text = [c.text.strip() for c in row.cells]
            if all(not t for t in cells_text):
                row._element.getparent().remove(row._element)
                empty_rows += 1
            else:
                break
        if empty_rows > 0:
            print(f"  Table: 已删除 {empty_rows} 个空行")

def remove_empty_paragraphs(doc):
    """删除文档中空段落（无文本且无特殊格式）"""
    removed = 0
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for p in list(doc.paragraphs):
        if not p.text.strip():
            # 保留有表格、分节符或图片的段落
            has_content = any(child.tag.endswith(('tbl', 'sectPr')) for child in p._element)
            if not has_content:
                # 检查 run 中是否有图片
                for r in p._element.findall(f'{{{ns}}}r'):
                    for pic_tag in ('drawing', 'pict', 'object'):
                        if r.find(f'{{{ns}}}{pic_tag}') is not None:
                            has_content = True
                            break
                    if has_content:
                        break
            if not has_content:
                p._element.getparent().remove(p._element)
                removed += 1
    if removed > 0:
        print(f"  已删除 {removed} 个空段落")

def _has_toc_field(doc):
    """检测文档是否已包含 TOC 域"""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for p in doc.paragraphs:
        # 检查 fldSimple
        fld_simple = p._element.find('.//w:fldSimple', ns)
        if fld_simple is not None:
            instr = fld_simple.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr')
            if instr and 'TOC' in instr:
                return True
        # 检查 fldChar
        fld_chars = p._element.findall('.//w:fldChar', ns)
        for fld in fld_chars:
            fld_type = fld.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
            if fld_type == 'begin':
                instr_texts = p._element.findall('.//w:instrText', ns)
                for it in instr_texts:
                    if it.text and 'TOC' in it.text:
                        return True
    return False

def _find_toc_paragraphs(tmpl_doc):
    """从模板中提取 TOC 域所在的段落 XML 元素列表"""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    toc_elements = []
    in_toc = False
    current_group = []

    for p in tmpl_doc.paragraphs:
        p_elem = p._element
        # 检测 TOC 开始
        fld_chars = p_elem.findall('.//w:fldChar', ns)
        for fld in fld_chars:
            fld_type = fld.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
            if fld_type == 'begin':
                instr_texts = p_elem.findall('.//w:instrText', ns)
                for it in instr_texts:
                    if it.text and 'TOC' in it.text:
                        in_toc = True
                        break
            elif fld_type == 'end':
                if in_toc:
                    current_group.append(copy.deepcopy(p_elem))
                    toc_elements.append(current_group)
                    current_group = []
                    in_toc = False

        # 检测 fldSimple 形式的 TOC
        fld_simple = p_elem.find('.//w:fldSimple', ns)
        if fld_simple is not None:
            instr = fld_simple.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr')
            if instr and 'TOC' in instr:
                toc_elements.append([copy.deepcopy(p_elem)])
                continue

        if in_toc:
            current_group.append(copy.deepcopy(p_elem))

    return toc_elements

def _insert_toc_after_heading(target_doc, toc_elements):
    """在目标文件的'目录'标题后插入 TOC 域段落"""
    # 找到'目录'段落的位置
    insert_after = None
    for i, p in enumerate(target_doc.paragraphs):
        if p.text.strip() == '目录':
            insert_after = p._element
            break

    if insert_after is None:
        print("  警告：未找到'目录'标题，跳过 TOC 插入")
        return False

    # 插入 TOC 段落
    for group in toc_elements:
        for elem in group:
            new_elem = copy.deepcopy(elem)
            insert_after.addnext(new_elem)
            next_elem = insert_after.getnext()
            if next_elem is None:
                print("  警告：TOC 域插入异常，终止")
                return False
            insert_after = next_elem

    return True

def fix_toc_from_template(tmpl_doc, target_doc):
    """
    同步目录域：如果目标文件没有 TOC 域，从模板复制。
    如果目标文件已有 TOC，仅报告不覆盖（避免破坏现有目录）。
    """
    if _has_toc_field(target_doc):
        print("  目标文件已包含 TOC 域，跳过")
        return

    toc_elements = _find_toc_paragraphs(tmpl_doc)
    if not toc_elements:
        print("  模板中未找到 TOC 域，跳过")
        return

    success = _insert_toc_after_heading(target_doc, toc_elements)
    if success:
        print(f"  已从模板复制 {len(toc_elements)} 组 TOC 域段落")
        print("  ⚠️  请在 Word 中右键目录 → '更新域' 刷新页码")

def fix_specific_issues(doc):
    """
    处理文件特定问题（根据实际需求修改）
    示例：删除多余段落、插入缺失标题、补全表格内容等
    """
    paras = doc.paragraphs

    # 示例：删除残留的多余段落
    # for p in paras:
    #     if '残留关键词' in p.text:
    #         p._element.getparent().remove(p._element)
    #         print("  已删除残留段落")
    #         break

    # 示例：插入缺失标题
    # for i, p in enumerate(paras):
    #     if '某正文内容' in p.text:
    #         if i > 0 and '缺失标题' not in paras[i-1].text:
    #             new_para = insert_paragraph_before(p, '（三）缺失标题')
    #             new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #             new_para.paragraph_format.line_spacing = LS_NORMAL
    #             new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    #             new_para.paragraph_format.space_before = Emu(SB_H1)
    #             new_para.paragraph_format.space_after = None
    #             new_para.paragraph_format.first_line_indent = Emu(FI_H2)
    #             for run in new_para.runs:
    #                 set_run_font(run, FONT_SONG, SIZE_BODY, True)
    #             print("  已插入缺失标题")
    #         break

    # 示例：补全表格最后一行
    # if len(doc.tables) > 6:
    #     table = doc.tables[6]
    #     last_row = table.rows[-1]
    #     cells = last_row.cells
    #     if len(cells) >= 6 and not cells[2].text.strip():
    #         cells[2].text = '设备名称'
    #         cells[3].text = '设备描述'
    #         cells[4].text = '数量'
    #         for p in cells[2].paragraphs + cells[3].paragraphs + cells[4].paragraphs:
    #             for run in p.runs:
    #                 set_run_font(run, FONT_SONG, SIZE_BODY, False)
    #         print("  已补全表格最后一行")

    pass  # 如无特定问题，保留空函数

def main():
    print(f"开始修复: {TARGET}")
    print(f"参考模板: {TEMPLATE}")
    print("-" * 60)

    doc = docx.Document(TARGET)

    # 尝试加载模板（用于自动同步隐藏格式）
    tmpl_doc = None
    try:
        tmpl_doc = docx.Document(TEMPLATE)
        print("✓ 模板加载成功，将自动同步隐藏格式")
    except Exception as e:
        print(f"⚠ 模板加载失败 ({e})，将使用 CONFIG 默认值")

    print("1. 修复段落格式...")
    classify_and_format(doc)

    print("2. 修复引号...")
    fix_quotes(doc)

    print("3. 修复表格字体...")
    fix_table_fonts(doc)

    print("4. 修复表格边框...")
    if tmpl_doc is not None:
        copy_table_borders_from_template(tmpl_doc, doc)
    else:
        fix_table_borders(doc)

    print("5. 修复表格对齐...")
    fix_table_alignment(doc)

    print("6. 同步页面设置...")
    fix_section_settings(doc, tmpl_doc)

    print("7. 同步目录域...")
    if tmpl_doc is not None:
        fix_toc_from_template(tmpl_doc, doc)
    else:
        print("  跳过（无模板）")

    print("8. 处理特定问题...")
    fix_specific_issues(doc)

    print("9. 删除空段落...")
    remove_empty_paragraphs(doc)

    print("10. 删除表格空行...")
    remove_empty_table_rows(doc)

    doc.save(TARGET)
    print("-" * 60)
    print(f"修复完成！已保存: {TARGET}")
    print("\n后续步骤:")
    print(f"  python3 verify_docx.py {TARGET} {TEMPLATE}")

if __name__ == '__main__':
    main()
