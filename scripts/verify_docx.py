#!/usr/bin/env python3
"""
多维度精确验证目标文件与模板的格式一致性
包含：段落、分节、页眉页脚、样式、表格
"""
import docx
import sys
import argparse

from utils import get_para_props, compare_para_props, check_file_exists
from utils import run_with_errors, check_write_permission, log_ok, log_warn, log_err

parser = argparse.ArgumentParser(description='多维度验证目标文件与模板的格式一致性')
parser.add_argument('target', nargs='?', default='target.docx', help='目标 docx 文件路径')
parser.add_argument('template', nargs='?', default='template.docx', help='模板 docx 文件路径')
args = parser.parse_args()
TARGET = args.target
TEMPLATE = args.template

# ============== 段落验证 ==============
def verify_paragraphs(tmpl_doc, target_doc):
    print("\n" + "=" * 90)
    print("【段落格式验证】")
    print("=" * 90)

    tmpl_map = {}
    for i, p in enumerate(tmpl_doc.paragraphs):
        info = get_para_props(p, max_text_len=70, with_outline=False)
        if info:
            tmpl_map[info['text']] = info

    diffs = []
    for i, p in enumerate(target_doc.paragraphs):
        info = get_para_props(p, max_text_len=70, with_outline=False)
        if not info:
            continue

        tmpl_info = None
        for t_text, t_info in tmpl_map.items():
            if info['text'][:15] == t_text[:15]:
                tmpl_info = t_info
                break

        if not tmpl_info:
            continue

        issues = compare_para_props(info, tmpl_info)

        if issues:
            diffs.append((i, info['text'], issues))

    for idx, text, issues in diffs[:50]:
        print(f"[{idx}] {text[:55]}")
        for iss in issues:
            print(f"    {iss}")

    print(f"\n段落验证: 共发现 {len(diffs)} 处差异")
    return len(diffs)

# ============== 分节验证 ==============
def verify_sections(tmpl_doc, target_doc):
    print("\n" + "=" * 90)
    print("【分节与页面设置验证】")
    print("=" * 90)

    diffs = []

    tmpl_sections = len(tmpl_doc.sections)
    target_sections = len(target_doc.sections)

    if tmpl_sections != target_sections:
        diffs.append(("分节数量", [f"分节数量不一致: 目标={target_sections}, 模板={tmpl_sections}"]))

    min_sections = min(tmpl_sections, target_sections)
    for i in range(min_sections):
        tmpl_s = tmpl_doc.sections[i]
        target_s = target_doc.sections[i]

        section_diffs = []
        if target_s.page_width != tmpl_s.page_width:
            section_diffs.append(f"page_width={target_s.page_width}(应{tmpl_s.page_width})")
        if target_s.page_height != tmpl_s.page_height:
            section_diffs.append(f"page_height={target_s.page_height}(应{tmpl_s.page_height})")
        if target_s.orientation != tmpl_s.orientation:
            section_diffs.append(f"orientation={target_s.orientation}(应{tmpl_s.orientation})")
        if target_s.top_margin != tmpl_s.top_margin:
            section_diffs.append(f"top_margin={target_s.top_margin}(应{tmpl_s.top_margin})")
        if target_s.bottom_margin != tmpl_s.bottom_margin:
            section_diffs.append(f"bottom_margin={target_s.bottom_margin}(应{tmpl_s.bottom_margin})")
        if target_s.left_margin != tmpl_s.left_margin:
            section_diffs.append(f"left_margin={target_s.left_margin}(应{tmpl_s.left_margin})")
        if target_s.right_margin != tmpl_s.right_margin:
            section_diffs.append(f"right_margin={target_s.right_margin}(应{tmpl_s.right_margin})")
        if target_s.header_distance != tmpl_s.header_distance:
            section_diffs.append(f"header_dist={target_s.header_distance}(应{tmpl_s.header_distance})")
        if target_s.footer_distance != tmpl_s.footer_distance:
            section_diffs.append(f"footer_dist={target_s.footer_distance}(应{tmpl_s.footer_distance})")
        if target_s.different_first_page_header_footer != tmpl_s.different_first_page_header_footer:
            section_diffs.append(f"diff_first_page={target_s.different_first_page_header_footer}(应{tmpl_s.different_first_page_header_footer})")

        if section_diffs:
            diffs.append((f"Section {i+1}", section_diffs))

    for section_name, issues in diffs:
        print(f"{section_name}:")
        for iss in issues:
            print(f"    {iss}")

    total = len(diffs)
    print(f"\n分节验证: 共发现 {total} 处差异")
    return total

# ============== 页眉页脚验证 ==============
def verify_headers_footers(tmpl_doc, target_doc):
    print("\n" + "=" * 90)
    print("【页眉页脚验证】")
    print("=" * 90)

    diffs = []
    min_sections = min(len(tmpl_doc.sections), len(target_doc.sections))

    for i in range(min_sections):
        tmpl_s = tmpl_doc.sections[i]
        target_s = target_doc.sections[i]

        # 页眉
        tmpl_header = tmpl_s.header
        target_header = target_s.header

        tmpl_header_text = '\n'.join([p.text for p in tmpl_header.paragraphs if p.text.strip()])
        target_header_text = '\n'.join([p.text for p in target_header.paragraphs if p.text.strip()])

        if tmpl_header_text != target_header_text:
            diffs.append((f"Section {i+1} 页眉", f"内容不一致"))

        # 页脚
        tmpl_footer = tmpl_s.footer
        target_footer = target_s.footer

        tmpl_footer_text = '\n'.join([p.text for p in tmpl_footer.paragraphs if p.text.strip()])
        target_footer_text = '\n'.join([p.text for p in target_footer.paragraphs if p.text.strip()])

        if tmpl_footer_text != target_footer_text:
            diffs.append((f"Section {i+1} 页脚", f"内容不一致"))

    for name, issue in diffs:
        print(f"{name}: {issue}")

    total = len(diffs)
    print(f"\n页眉页脚验证: 共发现 {total} 处差异")
    return total

# ============== 样式验证 ==============
def verify_styles(tmpl_doc, target_doc):
    print("\n" + "=" * 90)
    print("【样式定义验证】")
    print("=" * 90)

    diffs = []

    # 收集模板中的关键样式
    tmpl_styles = {}
    for style in tmpl_doc.styles:
        if style.type == 1:  # paragraph style
            key_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                          'toc 1', 'toc 2', 'toc 3', 'Header', 'Footer',
                          'Title', 'Subtitle']
            if style.name in key_styles:
                tmpl_styles[style.name] = style

    # 检查目标文件是否有这些样式
    target_style_names = {s.name for s in target_doc.styles}

    for name, tmpl_style in tmpl_styles.items():
        if name not in target_style_names:
            diffs.append(f"缺少样式: {name}")
            continue

        target_style = None
        for s in target_doc.styles:
            if s.name == name:
                target_style = s
                break

        if target_style and tmpl_style.font and target_style.font:
            if tmpl_style.font.name and target_style.font.name != tmpl_style.font.name:
                diffs.append(f"样式 '{name}' font={target_style.font.name}(应{tmpl_style.font.name})")
            if tmpl_style.font.size and target_style.font.size != tmpl_style.font.size:
                diffs.append(f"样式 '{name}' size={target_style.font.size}(应{tmpl_style.font.size})")
            if tmpl_style.font.bold is not None and target_style.font.bold != tmpl_style.font.bold:
                diffs.append(f"样式 '{name}' bold={target_style.font.bold}(应{tmpl_style.font.bold})")

    for issue in diffs[:30]:
        print(f"  {issue}")

    total = len(diffs)
    print(f"\n样式验证: 共发现 {total} 处差异")
    return total

# ============== 表格验证 ==============
def verify_tables(tmpl_doc, target_doc):
    print("\n" + "=" * 90)
    print("【表格验证】")
    print("=" * 90)

    diffs = []

    tmpl_tables = len(tmpl_doc.tables)
    target_tables = len(target_doc.tables)

    if tmpl_tables != target_tables:
        diffs.append(f"表格数量不一致: 目标={target_tables}, 模板={tmpl_tables}")

    min_tables = min(tmpl_tables, target_tables)
    for i in range(min_tables):
        tmpl_t = tmpl_doc.tables[i]
        target_t = target_doc.tables[i]

        tmpl_rows = len(tmpl_t.rows)
        target_rows = len(target_t.rows)
        tmpl_cols = len(tmpl_t.columns)
        target_cols = len(target_t.columns)

        if tmpl_rows != target_rows:
            diffs.append(f"Table {i+1} 行数不一致: 目标={target_rows}, 模板={tmpl_rows}")
        if tmpl_cols != target_cols:
            diffs.append(f"Table {i+1} 列数不一致: 目标={target_cols}, 模板={tmpl_cols}")

        # 对比单元格内容
        min_rows = min(tmpl_rows, target_rows)
        min_cols = min(tmpl_cols, target_cols)

        for r in range(min_rows):
            for c in range(min_cols):
                tmpl_text = tmpl_t.rows[r].cells[c].text.strip()
                target_text = target_t.rows[r].cells[c].text.strip()
                if tmpl_text != target_text:
                    diffs.append(f"Table {i+1}[{r},{c}] 内容不一致")

    for issue in diffs[:30]:
        print(f"  {issue}")

    total = len(diffs)
    print(f"\n表格验证: 共发现 {total} 处差异")
    return total

# ============== 主程序 ==============
@run_with_errors
def main():
    check_file_exists(TARGET, '目标文件')
    check_file_exists(TEMPLATE, '模板文件')

    tmpl = docx.Document(TEMPLATE)
    doc = docx.Document(TARGET)

    print("=" * 90)
    print(f"多维度验证: {TARGET} vs {TEMPLATE}")
    print("=" * 90)

    total_diffs = 0
    total_diffs += verify_paragraphs(tmpl, doc)
    total_diffs += verify_sections(tmpl, doc)
    total_diffs += verify_headers_footers(tmpl, doc)
    total_diffs += verify_styles(tmpl, doc)
    total_diffs += verify_tables(tmpl, doc)

    print("\n" + "=" * 90)
    if total_diffs == 0:
        log_ok("全部验证通过！目标文件与模板格式完全一致。")
    else:
        log_warn(f"验证未通过！共发现 {total_diffs} 处差异，请修复后重新验证。")
    print("=" * 90)

if __name__ == '__main__':
    main()
