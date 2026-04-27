#!/usr/bin/env python3
"""
将模板文件中的页眉页脚复制到目标文件
⚠️ 警告：这会覆盖目标文件的所有页眉页脚内容！
"""
import docx
from copy import deepcopy
import sys
import argparse
from utils import get_odd_even_headers, set_odd_even_headers, check_file_exists
from utils import run_with_errors, check_write_permission, log_ok, log_warn, log_err

parser = argparse.ArgumentParser(description='将模板页眉页脚复制到目标文件')
parser.add_argument('template', nargs='?', default='template.docx', help='模板 docx 文件路径')
parser.add_argument('target', nargs='?', default='target.docx', help='目标 docx 文件路径')
args = parser.parse_args()
TMPL = args.template
TARGET = args.target

def copy_headers_footers(tmpl_doc, target_doc):
    """复制页眉页脚"""

    min_sections = min(len(tmpl_doc.sections), len(target_doc.sections))

    print("页眉页脚复制报告:")
    print("-" * 60)

    for i in range(min_sections):
        tmpl_s = tmpl_doc.sections[i]
        target_s = target_doc.sections[i]

        # 复制页眉
        tmpl_header = tmpl_s.header
        target_header = target_s.header

        if tmpl_header.paragraphs or tmpl_header.tables:
            # 清空目标页眉
            for p in list(target_header.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(target_header.tables):
                t._element.getparent().remove(t._element)

            # 复制模板页眉内容
            for element in tmpl_header._element:
                if element.tag.endswith(('p', 'tbl')):
                    target_header._element.append(deepcopy(element))
            print(f"  [复制] Section {i+1} 页眉")

        # 复制页脚
        tmpl_footer = tmpl_s.footer
        target_footer = target_s.footer

        if tmpl_footer.paragraphs or tmpl_footer.tables:
            # 清空目标页脚
            for p in list(target_footer.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(target_footer.tables):
                t._element.getparent().remove(t._element)

            # 复制模板页脚内容
            for element in tmpl_footer._element:
                if element.tag.endswith(('p', 'tbl')):
                    target_footer._element.append(deepcopy(element))
            print(f"  [复制] Section {i+1} 页脚")

        # 复制页面设置属性
        target_s.different_first_page_header_footer = tmpl_s.different_first_page_header_footer
        set_odd_even_headers(target_s, get_odd_even_headers(tmpl_s))

    print("-" * 60)
    print("页眉页脚复制完成！")

@run_with_errors
def main():
    check_file_exists(TMPL, '模板文件')
    check_file_exists(TARGET, '目标文件')
    check_write_permission(TARGET, '目标文件')

    tmpl = docx.Document(TMPL)
    doc = docx.Document(TARGET)

    copy_headers_footers(tmpl, doc)
    doc.save(TARGET)
    log_ok(f"已保存: {TARGET}")

if __name__ == '__main__':
    main()
