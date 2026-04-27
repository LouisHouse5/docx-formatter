#!/usr/bin/env python3
"""
全面对比目标文件与模板的差异（显式格式 + 隐藏格式）
包含：段落、分节、页眉页脚、样式、表格、目录
"""
import docx
import sys
import argparse
from utils import get_para_props, compare_para_props, check_file_exists
from utils import run_with_errors, check_write_permission, log_ok, log_warn, log_err

parser = argparse.ArgumentParser(description='全面对比目标文件与模板的差异')
parser.add_argument('target', nargs='?', default='target.docx', help='目标 docx 文件路径')
parser.add_argument('template', nargs='?', default='template.docx', help='模板 docx 文件路径')
args = parser.parse_args()
TARGET = args.target
TEMPLATE = args.template

check_file_exists(TARGET, '目标文件')
check_file_exists(TEMPLATE, '模板文件')

print("=" * 80)
print(f"目标文件: {TARGET}")
print(f"模板文件: {TEMPLATE}")
print("=" * 80)

tmpl = docx.Document(TEMPLATE)
doc = docx.Document(TARGET)

# ============== 1. 段落格式差异 ==============
print("\n【1. 段落格式差异】")

tmpl_map = {}
for i, p in enumerate(tmpl.paragraphs):
    info = get_para_props(p)
    if info:
        tmpl_map[info['text'][:15]] = info

para_diffs = []
for i, p in enumerate(doc.paragraphs):
    info = get_para_props(p)
    if not info:
        continue

    issues = []
    if info['font_name'] is None:
        issues.append('font=None')
    if info['font_size'] is None:
        issues.append('size=None')
    if issues:
        print(f"  [{i:3}] {info['text'][:50]:<50} -> {', '.join(issues)}")

    # 对比模板
    tmpl_info = None
    for t_text, t_info in tmpl_map.items():
        if info['text'][:15] == t_text[:15]:
            tmpl_info = t_info
            break

    if not tmpl_info:
        continue

    comp_issues = compare_para_props(info, tmpl_info)

    if comp_issues:
        para_diffs.append((i, info['text'], comp_issues))

for idx, text, issues in para_diffs[:30]:
    print(f"  [{idx:3}] {text[:50]}")
    for iss in issues:
        print(f"      {iss}")

print(f"  段落差异: {len(para_diffs)} 处")

# ============== 2. 分节差异 ==============
print("\n【2. 分节与页面设置差异】")
section_diffs = []

tmpl_sections = len(tmpl.sections)
target_sections = len(doc.sections)
if tmpl_sections != target_sections:
    section_diffs.append(f"分节数量: 目标={target_sections}, 模板={tmpl_sections}")

min_sections = min(tmpl_sections, target_sections)
for i in range(min_sections):
    tmpl_s = tmpl.sections[i]
    target_s = doc.sections[i]

    if target_s.page_width != tmpl_s.page_width:
        section_diffs.append(f"Section {i+1} page_width: {target_s.page_width} != {tmpl_s.page_width}")
    if target_s.page_height != tmpl_s.page_height:
        section_diffs.append(f"Section {i+1} page_height: {target_s.page_height} != {tmpl_s.page_height}")
    if target_s.orientation != tmpl_s.orientation:
        section_diffs.append(f"Section {i+1} orientation: {target_s.orientation} != {tmpl_s.orientation}")
    if target_s.top_margin != tmpl_s.top_margin:
        section_diffs.append(f"Section {i+1} top_margin: {target_s.top_margin} != {tmpl_s.top_margin}")
    if target_s.bottom_margin != tmpl_s.bottom_margin:
        section_diffs.append(f"Section {i+1} bottom_margin: {target_s.bottom_margin} != {tmpl_s.bottom_margin}")
    if target_s.left_margin != tmpl_s.left_margin:
        section_diffs.append(f"Section {i+1} left_margin: {target_s.left_margin} != {tmpl_s.left_margin}")
    if target_s.right_margin != tmpl_s.right_margin:
        section_diffs.append(f"Section {i+1} right_margin: {target_s.right_margin} != {tmpl_s.right_margin}")
    if target_s.header_distance != tmpl_s.header_distance:
        section_diffs.append(f"Section {i+1} header_distance: {target_s.header_distance} != {tmpl_s.header_distance}")
    if target_s.footer_distance != tmpl_s.footer_distance:
        section_diffs.append(f"Section {i+1} footer_distance: {target_s.footer_distance} != {tmpl_s.footer_distance}")
    if target_s.different_first_page_header_footer != tmpl_s.different_first_page_header_footer:
        section_diffs.append(f"Section {i+1} diff_first_page: {target_s.different_first_page_header_footer} != {tmpl_s.different_first_page_header_footer}")

for issue in section_diffs:
    print(f"  {issue}")

if not section_diffs:
    print("  无差异")
print(f"  分节差异: {len(section_diffs)} 处")

# ============== 3. 页眉页脚差异 ==============
print("\n【3. 页眉页脚差异】")
hf_diffs = []
min_sections = min(len(tmpl.sections), len(doc.sections))

for i in range(min_sections):
    tmpl_s = tmpl.sections[i]
    target_s = doc.sections[i]

    tmpl_header = '\n'.join([p.text for p in tmpl_s.header.paragraphs if p.text.strip()])
    target_header = '\n'.join([p.text for p in target_s.header.paragraphs if p.text.strip()])
    if tmpl_header != target_header:
        hf_diffs.append(f"Section {i+1} 页眉内容不一致")

    tmpl_footer = '\n'.join([p.text for p in tmpl_s.footer.paragraphs if p.text.strip()])
    target_footer = '\n'.join([p.text for p in target_s.footer.paragraphs if p.text.strip()])
    if tmpl_footer != target_footer:
        hf_diffs.append(f"Section {i+1} 页脚内容不一致")

for issue in hf_diffs:
    print(f"  {issue}")

if not hf_diffs:
    print("  无差异")
print(f"  页眉页脚差异: {len(hf_diffs)} 处")

# ============== 4. 样式差异 ==============
print("\n【4. 样式定义差异】")
style_diffs = []

tmpl_styles = {}
for style in tmpl.styles:
    if style.type == 1:  # paragraph style
        key_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                      'toc 1', 'toc 2', 'toc 3', 'Header', 'Footer']
        if style.name in key_styles:
            tmpl_styles[style.name] = style

target_style_names = {s.name for s in doc.styles}

for name, tmpl_style in tmpl_styles.items():
    if name not in target_style_names:
        style_diffs.append(f"缺少样式: {name}")
        continue

    target_style = None
    for s in doc.styles:
        if s.name == name:
            target_style = s
            break

    if target_style and tmpl_style.font and target_style.font:
        if tmpl_style.font.name and target_style.font.name != tmpl_style.font.name:
            style_diffs.append(f"样式 '{name}' font: {target_style.font.name} != {tmpl_style.font.name}")
        if tmpl_style.font.size and target_style.font.size != tmpl_style.font.size:
            style_diffs.append(f"样式 '{name}' size: {target_style.font.size} != {tmpl_style.font.size}")

for issue in style_diffs:
    print(f"  {issue}")

if not style_diffs:
    print("  无差异")
print(f"  样式差异: {len(style_diffs)} 处")

# ============== 5. 表格差异 ==============
print("\n【5. 表格差异】")
table_diffs = []

tmpl_tables = len(tmpl.tables)
target_tables = len(doc.tables)
if tmpl_tables != target_tables:
    table_diffs.append(f"表格数量: 目标={target_tables}, 模板={tmpl_tables}")

min_tables = min(tmpl_tables, target_tables)
for i in range(min_tables):
    tmpl_t = tmpl.tables[i]
    target_t = doc.tables[i]

    if len(tmpl_t.rows) != len(target_t.rows):
        table_diffs.append(f"Table {i+1} 行数: {len(target_t.rows)} != {len(tmpl_t.rows)}")
    if len(tmpl_t.columns) != len(target_t.columns):
        table_diffs.append(f"Table {i+1} 列数: {len(target_t.columns)} != {len(tmpl_t.columns)}")

    min_rows = min(len(tmpl_t.rows), len(target_t.rows))
    min_cols = min(len(tmpl_t.columns), len(target_t.columns))
    for r in range(min_rows):
        for c in range(min_cols):
            if tmpl_t.rows[r].cells[c].text.strip() != target_t.rows[r].cells[c].text.strip():
                table_diffs.append(f"Table {i+1}[{r},{c}] 内容不一致")

for issue in table_diffs[:20]:
    print(f"  {issue}")

if len(table_diffs) > 20:
    print(f"  ... 还有 {len(table_diffs)-20} 处差异")
print(f"  表格差异: {len(table_diffs)} 处")

# ============== 汇总 ==============
total = len(para_diffs) + len(section_diffs) + len(hf_diffs) + len(style_diffs) + len(table_diffs)
print("\n" + "=" * 80)
if total == 0:
    log_ok("审核汇总: 共发现 0 处差异，全部通过！")
else:
    log_warn(f"审核汇总: 共发现 {total} 处差异")
print(f"  - 段落格式: {len(para_diffs)}")
print(f"  - 分节设置: {len(section_diffs)}")
print(f"  - 页眉页脚: {len(hf_diffs)}")
print(f"  - 样式定义: {len(style_diffs)}")
print(f"  - 表格: {len(table_diffs)}")
print("=" * 80)
